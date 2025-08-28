
import json
import os
import jieba
import numpy as np
from openai import OpenAI
from bm25 import BM25
from langchain_core.documents import Document
'''
基于RAG来介绍淘宝相关规则
用bm25做召回
同样以ALI的api作为我们的大模型
https://docs.bigmodel.cn/cn/guide/start/model-overview
'''

#ALI的api作为我们的大模型
def call_large_model(prompt):
    client = OpenAI(api_key=os.environ.get("ALI_API_KEY"),
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",) # 填写您自己的APIKey
    response = client.chat.completions.create(
        model="qwen-plus",  # 填写需要调用的模型名称
        messages=[
            {"role": "user", "content": prompt},
        ],
    )
    response_text = response.choices[0].message.content
    return response_text


class TaobaoRulesRAG:
    def __init__(self, docs: list[Document]):
        """
        初始化RAG系统
        :param docs: 从docx提取的内容块列表（每个元素是LangChain的Document对象）
        """
        self.load_rules_data(docs)  # 加载内容块数据

    def load_rules_data(self, docs: list[Document]):
        # 存储原始内容块（key: 索引, value: 文本内容）
        self.rules_data = {i: doc.page_content for i, doc in enumerate(docs)}

        # 构建BM25语料库（分词处理）
        self.corpus = {}  # key: 内容块索引, value: 分词列表
        self.index_to_content = {}  # 索引到原始文本的映射（和rules_data一致，便于理解）

        for idx, doc in enumerate(docs):
            text = doc.page_content
            self.corpus[idx] = jieba.lcut(text)  # 分词
            self.index_to_content[idx] = text  # 保存原始文本

        # 初始化BM25模型
        self.bm25_model = BM25(self.corpus)
        print(f"成功加载 {len(docs)} 个内容块到RAG系统")

    def retrieve(self, user_query: str) -> str:
        """根据用户查询，用BM25召回最相关的内容块"""
        # 对查询进行分词
        query_tokens = jieba.lcut(user_query)
        # 获取所有内容块的BM25分数
        scores = self.bm25_model.get_scores(query_tokens)  # 假设返回格式为[(索引, 分数), ...]
        # 按分数降序排序，取最相关的内容块
        sorted_scores = sorted(scores, key=lambda x: x[1], reverse=True)
        if not sorted_scores:
            return "未找到相关规则"
        # 获取最相关内容块的原始文本
        top_idx = sorted_scores[0][0]
        return self.index_to_content[top_idx]

    def query(self, user_query: str):
        """完整流程：检索→生成回答"""
        print(f"用户查询：{user_query}")
        print("=" * 50)

        # 1. 检索相关内容块
        retrieved_text = self.retrieve(user_query)
        print(f"检索到的规则内容：\n{retrieved_text}")
        print("=" * 50)

        # 2. 构造提示词，调用大模型
        prompt = f"""
        请根据以下淘宝规则内容，准确回答用户问题。
        规则内容：{retrieved_text}
        用户问题：{user_query}
        回答需简洁明了，直接引用规则内容，不添加无关信息。
        """
        response = call_large_model(prompt)
        print(f"模型回答：\n{response}")
        print("=" * 50)


# 主函数：加载内容块并测试
if __name__ == "__main__":
    # 1. 先加载之前提取的淘宝规则内容块（这里复用你之前的加载函数）
    from docx import Document as DocxDocument


    def load_docx_with_tables(file_path):
        """复用提取内容块的函数"""
        docx_doc = DocxDocument(file_path)
        docs = []
        # 提取段落
        for para in docx_doc.paragraphs:
            if para.text.strip():
                docs.append(Document(
                    page_content=para.text,
                    metadata={"type": "Text"}
                ))
        # 提取表格（用|分隔单元格）
        for table in docx_doc.tables:
            table_text = []
            for row in table.rows:
                row_text = "|".join([cell.text.strip() for cell in row.cells])
                table_text.append(row_text)
            docs.append(Document(
                page_content="\n".join(table_text),
                metadata={"type": "Table"}
            ))
        return docs


    # 2. 加载你的docx文件（替换为实际路径）
    docx_path = "D:/PycharmProjects/AI学习预习/week14+大语言模型相关第四讲/week14/淘宝网超时说明.docx"
    rules_docs = load_docx_with_tables(docx_path)

    # 3. 初始化RAG系统并测试查询
    rag = TaobaoRulesRAG(rules_docs)

    # 测试查询（替换为实际的淘宝规则问题）
    test_queries = [
        "买家已付款后，卖家多久不发货会关闭交易？",
        "等待买家付款的状态下，超时时间是多少？",
        "虚假交易的处罚措施是什么？"
    ]

    for query in test_queries:
        rag.query(query)
        print("\n" + "-" * 80 + "\n")
        print("----------------")
        print("No RAG (直接请求大模型回答)：")
        print(call_large_model(query))

# if __name__ == "__main__":
#     rag = SimpleRAG()
#     user_query = "高射火炮是谁的技能"
#     rag.query(user_query)
#
#     print("----------------")
#     print("No RAG (直接请求大模型回答)：")
#     print(call_large_model(user_query))
